import json
from pathlib import Path

import pytest

from officekit import create_document, extract_text, ops, view, batch, importer
from officekit.cli import main


# ----------------------------------------------------------------------
# view: outline / stats / issues
# ----------------------------------------------------------------------
def test_view_outline_docx(tmp_path):
    p = create_document("docx", tmp_path / "d.docx", title="T")
    from docx import Document as DocxDocument

    doc = DocxDocument(str(p))
    doc.add_heading("Chapter 1", level=1)
    doc.save(str(p))
    out = view.view_outline(p)
    assert out["kind"] == "docx"
    assert any(h["text"] == "Chapter 1" for h in out["headings"])


def test_view_stats_pptx(tmp_path):
    p = create_document("pptx", tmp_path / "d.pptx", title="T")
    s = view.view_stats(p)
    assert s["kind"] == "pptx" and s["slides"] == 1


def test_view_issues_flags_unfilled_placeholder(tmp_path):
    from docx import Document as DocxDocument

    p = tmp_path / "d.docx"
    doc = DocxDocument()
    doc.add_paragraph("Hi {{x}}")
    doc.save(str(p))
    issues = view.view_issues(p)
    assert any(i["code"] == "unfilled_placeholder" for i in issues)
def test_docx_get_set_roundtrip(tmp_path):
    p = create_document("docx", tmp_path / "d.docx", title="T")
    ops.set_value(p, "/paragraphs[1]", "Edited")
    assert ops.get(p, "/paragraphs[1]") == "Edited"
    assert "Edited" in extract_text(p)


def test_pptx_set_title_and_add_slide(tmp_path):
    p = create_document("pptx", tmp_path / "d.pptx", title="T")
    ops.set_value(p, "/slides[1]/shapes[1]", "NewTitle")
    assert ops.get(p, "/slides[1]/shapes[1]") == "NewTitle"
    ops.add(p, "slide", value="Second")
    assert view.view_stats(p)["slides"] == 2


def test_xlsx_set_get_add_row(tmp_path):
    p = create_document("xlsx", tmp_path / "b.xlsx", title="S")
    ops.set_value(p, "/sheets[1]/A1", "hello")
    assert ops.get(p, "/sheets[1]/A1") == "hello"
    ops.add(p, "row", value="1\t2\t3")
    from openpyxl import load_workbook

    ws = load_workbook(str(p)).active
    assert list(ws.iter_rows(min_row=ws.max_row, max_row=ws.max_row, values_only=True))[
        0
    ] == ("1", "2", "3")


def test_remove_docx_paragraph(tmp_path):
    p = create_document("docx", tmp_path / "d.docx", title="T")
    ops.add(p, "paragraph", value="toast")
    ops.remove(p, "/paragraphs[3]")  # heading[1], "Created..."[2], toast[3]
    assert "toast" not in extract_text(p)


# ----------------------------------------------------------------------
# batch: atomic apply + rollback
# ----------------------------------------------------------------------
def test_batch_applies_all(tmp_path):
    p = create_document("xlsx", tmp_path / "b.xlsx", title="S")
    ops.set_value(p, "/sheets[1]/A1", "orig")
    batch.run_batch(p, [{"op": "set", "target": "/sheets[1]/A1", "value": "new"}])
    assert ops.get(p, "/sheets[1]/A1") == "new"


def test_batch_rolls_back_on_failure(tmp_path):
    p = create_document("xlsx", tmp_path / "b.xlsx", title="S")
    ops.set_value(p, "/sheets[1]/A1", "orig")
    with pytest.raises(Exception):
        batch.run_batch(
            p,
            [
                {"op": "set", "target": "/sheets[1]/A1", "value": "new"},
                {"op": "remove", "target": "/sheets[99]"},
            ],
        )
    assert ops.get(p, "/sheets[1]/A1") == "orig"


# ----------------------------------------------------------------------
# import_csv
# ----------------------------------------------------------------------
def test_import_csv_creates_worksheet(tmp_path):
    csv_path = tmp_path / "data.csv"
    csv_path.write_text("name,age\nAlice,30\nBob,25", encoding="utf-8")
    out = tmp_path / "book.xlsx"
    importer.import_csv(out, "People", csv_path, header=True)
    from openpyxl import load_workbook

    wb = load_workbook(str(out))
    assert "People" in wb.sheetnames
    assert wb["People"]["A2"].value == "Alice"


# ----------------------------------------------------------------------
# CLI --json structured output
# ----------------------------------------------------------------------
def test_cli_json_success(tmp_path, capsys):
    p = create_document("docx", tmp_path / "d.docx", title="T")
    assert main(["info", str(p), "--json"]) == 0
    out = json.loads(capsys.readouterr().out)
    assert out["success"] is True and out["data"]["kind"] == "docx"


def test_cli_json_error_is_structured(tmp_path, capsys):
    assert main(["info", str(tmp_path / "missing.docx"), "--json"]) == 1
    out = json.loads(capsys.readouterr().out)
    assert out["success"] is False and "error" in out


def test_cli_view_mode_and_batch(tmp_path, capsys):
    p = create_document("xlsx", tmp_path / "b.xlsx", title="S")
    assert main(["view", str(p), "--mode", "stats", "--json"]) == 0
    stats = json.loads(capsys.readouterr().out)["data"]
    assert stats["kind"] == "xlsx"
    ops.set_value(p, "/sheets[1]/A1", "x")
    assert (
        main(
            [
                "batch",
                str(p),
                "--operations",
                '[{"op":"set","target":"/sheets[1]/A1","value":"y"}]',
            ]
        )
        == 0
    )
    assert ops.get(p, "/sheets[1]/A1") == "y"


def test_cli_import(tmp_path):
    csv_path = tmp_path / "data.csv"
    csv_path.write_text("a,b\n1,2", encoding="utf-8")
    out = tmp_path / "book.xlsx"
    assert main(["import", str(out), "Sheet1", str(csv_path), "--header"]) == 0
    from openpyxl import load_workbook

    assert load_workbook(str(out))["Sheet1"]["A1"].value == "a"
