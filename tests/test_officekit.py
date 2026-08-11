from pathlib import Path

import pytest
from docx import Document as DocxDocument
from officekit import (
    create_document,
    document_info,
    extract_text,
    merge_template,
    validate_document,
)
from officekit.cli import main


@pytest.mark.parametrize("kind", ["docx", "xlsx", "pptx"])
def test_create_validate_info(kind, tmp_path):
    p = create_document(kind, tmp_path / f"sample.{kind}", title="Demo")
    assert p.exists()
    assert validate_document(p) == []
    info = document_info(p)
    assert info["kind"] == kind


def test_extract_text_docx(tmp_path):
    p = tmp_path / "doc.docx"
    doc = DocxDocument()
    doc.add_paragraph("Hello world")
    doc.save(str(p))
    assert "Hello world" in extract_text(p)


def test_merge_docx_fills_placeholder(tmp_path):
    p = tmp_path / "letter.docx"
    doc = DocxDocument()
    doc.add_paragraph("Dear {{name}},")
    doc.save(str(p))

    out = merge_template(p, tmp_path / "out.docx", {"name": "Ada"})
    assert "Dear Ada," in extract_text(out)


def test_merge_docx_keeps_unknown_placeholder(tmp_path):
    p = tmp_path / "t.docx"
    doc = DocxDocument()
    doc.add_paragraph("Hi {{known}} and {{missing}}")
    doc.save(str(p))

    out = merge_template(p, tmp_path / "o.docx", {"known": "X"})
    text = extract_text(out)
    assert "Hi X and {{missing}}" in text


def test_merge_xlsx_fills_cell(tmp_path):
    from openpyxl import Workbook

    p = tmp_path / "book.xlsx"
    wb = Workbook()
    wb.active["A1"] = "Hello {{name}}"
    wb.save(str(p))

    out = merge_template(p, tmp_path / "out.xlsx", {"name": "Bob"})
    from openpyxl import load_workbook

    val = load_workbook(str(out)).active["A1"].value
    assert val == "Hello Bob"


def test_merge_pptx_fills_title(tmp_path):
    p = create_document("pptx", tmp_path / "deck.pptx", title="Hi {{who}}")
    out = merge_template(p, tmp_path / "out.pptx", {"who": "Team"})
    assert "Hi Team" in extract_text(out)


def test_cli_text_and_info(tmp_path, capsys):
    # Build a docx we can read back through the CLI
    p = tmp_path / "d.docx"
    doc = DocxDocument()
    doc.add_paragraph("Readable")
    doc.save(str(p))

    assert main(["text", str(p)]) == 0
    out = capsys.readouterr().out
    assert "Readable" in out

    capsys.readouterr()
    assert main(["info", str(p)]) == 0
    assert "docx" in capsys.readouterr().out


def test_cli_validate_failure(tmp_path):
    missing = tmp_path / "nope.docx"
    assert main(["validate", str(missing)]) == 1


def test_detect_kind_rejects_unknown(tmp_path):
    with pytest.raises(ValueError):
        extract_text(tmp_path / "x.txt")
